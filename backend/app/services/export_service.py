import os
import uuid
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx import Document


class ExportService:
    @staticmethod
    def export_to_word(report):
        filename = f"report_{uuid.uuid4().hex}.docx"
        filepath = os.path.join("tmp", filename)

        os.makedirs("tmp", exist_ok=True)

        doc = Document()
        doc.add_heading(report.title, 0)

        for line in report.full_content.split("\n"):
            doc.add_paragraph(line)

        doc.save(filepath)
        return filepath

    @staticmethod
    def export_to_pdf(report):
        filename = f"report_{uuid.uuid4().hex}.pdf"
        filepath = os.path.join("tmp", filename)

        os.makedirs("tmp", exist_ok=True)

        c = canvas.Canvas(filepath, pagesize=A4)
        width, height = A4

        x = 50
        y = height - 50

        c.setFont("Times-Bold", 16)
        c.drawString(x, y, report.title)
        y -= 40

        c.setFont("Times-Roman", 11)

        for line in report.full_content.split("\n"):
            if y < 50:
                c.showPage()
                c.setFont("Times-Roman", 11)
                y = height - 50

            c.drawString(x, y, line[:1000])
            y -= 14

        c.save()
        return filepath
